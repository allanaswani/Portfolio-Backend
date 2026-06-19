"""Client Brief API — CRUD plus faithful Word-document generation.

``build_brief_docx`` regenerates the HFCB client-brief memo in the exact layout
of the signed template (title heading, 3x2 memo table, BACKGROUND narrative,
bulleted responsibilities, numbered business opportunities, footer) so a director
can download a ready-to-sign ``.docx`` for any client an RM has captured.
"""

from io import BytesIO

import django_filters.rest_framework
from django.http import HttpResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from drf_spectacular.utils import extend_schema
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView

from core.pagination import StandardPagination
from .models import ClientBrief
from .serializers import ClientBriefSerializer

DjangoFilterBackend = django_filters.rest_framework.DjangoFilterBackend
TAG = ["Client Briefs"]

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ── Document builder ─────────────────────────────────────────────────────────

def _paragraphs(text):
    """Split a text block into paragraphs on blank lines (trims empties)."""
    if not text:
        return []
    blocks = str(text).replace("\r\n", "\n").split("\n\n")
    return [b.strip() for b in blocks if b.strip()]


def _lines(text):
    """Non-empty lines of a text block, preserving order."""
    if not text:
        return []
    return [ln.strip() for ln in str(text).replace("\r\n", "\n").split("\n") if ln.strip()]


def _bold_para(doc, text, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    return p


def _cell_label(cell, label, value):
    """Render a memo-table cell as a bold label followed by its value."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(label)
    run.bold = True
    if value:
        para.add_run(f" {value}")


def build_brief_docx(brief):
    """Return a ``BytesIO`` holding the brief rendered as a .docx memo."""
    doc = Document()

    # Page geometry + base font to match the signed template (Letter, 12pt).
    section = doc.sections[0]
    section.page_width = Pt(612)      # 8.5"
    section.page_height = Pt(792)     # 11"
    section.left_margin = Pt(57.6)    # 0.8"
    section.right_margin = Pt(45)     # 0.625"
    section.top_margin = Pt(36)       # 0.5"
    section.bottom_margin = Pt(36)
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(12)

    client = brief.client_name.strip()

    # Title — same built-in Heading 1 style as the template, centred.
    title = doc.add_heading(f"HFCB CLIENT BRIEF – {client}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Memo header table (3 rows x 2 cols, bordered).
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    _cell_label(table.rows[0].cells[0], "FROM:", brief.from_party)
    _cell_label(table.rows[0].cells[1], "SUBJECT:", brief.display_subject)
    _cell_label(table.rows[1].cells[0], "TO:", brief.to_party)
    _cell_label(table.rows[1].cells[1], "DATE:", brief.brief_date.strftime("%d/%m/%Y"))
    _cell_label(table.rows[2].cells[0], f"CONTACT PERSONS – {client}", "")
    _cell_label(table.rows[2].cells[1], "", brief.contact_persons)

    doc.add_paragraph()

    # BACKGROUND
    _bold_para(doc, "BACKGROUND")
    if client:
        _bold_para(doc, client)
    for para in _paragraphs(brief.background):
        doc.add_paragraph(para)

    # Key responsibilities (bulleted)
    responsibilities = _lines(brief.key_responsibilities)
    if responsibilities:
        _bold_para(doc, "Key responsibilities include:", size=12)
        for item in responsibilities:
            doc.add_paragraph(item, style="List Bullet")

    if brief.budget_note.strip():
        for para in _paragraphs(brief.budget_note):
            doc.add_paragraph(para)

    # Business opportunities (numbered headings + body)
    opportunities = [o for o in (brief.opportunities or []) if o.get("title") or o.get("body")]
    if opportunities:
        doc.add_paragraph()
        _bold_para(doc, "Business Opportunities")
        for opp in opportunities:
            title_text = str(opp.get("title", "")).strip()
            if title_text:
                _bold_para(doc, title_text)
            for line in _lines(opp.get("body", "")):
                doc.add_paragraph(line)

    # Footer
    doc.add_paragraph()
    if brief.prepared_by.strip():
        doc.add_paragraph(f"Prepared by: {brief.prepared_by.strip()}")
    if brief.rm_designation.strip():
        doc.add_paragraph(brief.rm_designation.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── CRUD ─────────────────────────────────────────────────────────────────────

@extend_schema(tags=TAG)
class ClientBriefListCreateView(generics.ListCreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ClientBriefSerializer
    pagination_class = StandardPagination
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ["status", "rm", "client_name"]
    queryset = ClientBrief.objects.select_related("rm").all()

    def perform_create(self, serializer):
        user = self.request.user
        prepared = serializer.validated_data.get("prepared_by", "").strip()
        if not prepared:
            prepared = user.get_full_name().strip() or user.username
        serializer.save(rm=user, prepared_by=prepared)


@extend_schema(tags=TAG)
class ClientBriefDetailView(generics.RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ClientBriefSerializer
    queryset = ClientBrief.objects.select_related("rm").all()


@extend_schema(tags=TAG, responses={(200, DOCX_CONTENT_TYPE): bytes})
class ClientBriefDownloadView(APIView):
    """GET briefs/<pk>/download/ — the signable Word memo for this brief."""

    permission_classes = [IsAuthenticated]

    def get(self, request, pk):
        brief = generics.get_object_or_404(ClientBrief, pk=pk)
        buf = build_brief_docx(brief)
        safe_name = brief.client_name.strip() or "CLIENT"
        filename = f"{safe_name} CLIENT BRIEF.docx"
        response = HttpResponse(buf.getvalue(), content_type=DOCX_CONTENT_TYPE)
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
