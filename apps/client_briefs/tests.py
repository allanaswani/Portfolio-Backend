from io import BytesIO

from django.contrib.auth.models import User
from docx import Document
from rest_framework.test import APITestCase

from .models import ClientBrief


class ClientBriefTests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="rm1", password="pw", first_name="Judith", last_name="Ireri"
        )
        self.client.force_authenticate(self.user)

    def _payload(self, **over):
        data = {
            "client_name": "ATHI WATER WORKS DEVELOPMENT AGENCY",
            "brief_date": "2026-06-12",
            "contact_persons": "Eng. Joseph Kamau - CEO",
            "background": "AWWDA is one of nine agencies.\n\nIt manages water infra.",
            "key_responsibilities": "Develop waterworks\nOperate the waterworks",
            "budget_note": "Annual budget ~564.6M.",
            "opportunities": [
                {"title": "i). Institutional Corporate Account", "body": "Open project account."},
                {"title": "ii). Staff Loans Scheme", "body": "Mortgage & car loans."},
            ],
        }
        data.update(over)
        return data

    def test_create_sets_rm_and_prepared_by(self):
        res = self.client.post("/client_briefs/briefs/", self._payload(), format="json")
        self.assertEqual(res.status_code, 201, res.content)
        brief = ClientBrief.objects.get()
        self.assertEqual(brief.rm, self.user)
        # prepared_by defaults to the RM's full name
        self.assertEqual(brief.prepared_by, "Judith Ireri")
        self.assertEqual(brief.status, "draft")
        self.assertEqual(len(brief.opportunities), 2)

    def test_display_subject_defaults(self):
        res = self.client.post("/client_briefs/briefs/", self._payload(), format="json")
        self.assertEqual(
            res.json()["display_subject"],
            "CLIENT BRIEF – ATHI WATER WORKS DEVELOPMENT AGENCY",
        )

    def test_list(self):
        self.client.post("/client_briefs/briefs/", self._payload(), format="json")
        res = self.client.get("/client_briefs/briefs/")
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.json()["count"], 1)

    def test_download_returns_valid_docx(self):
        create = self.client.post("/client_briefs/briefs/", self._payload(), format="json")
        pk = create.json()["id"]
        res = self.client.get(f"/client_briefs/briefs/{pk}/download/")
        self.assertEqual(res.status_code, 200)
        self.assertEqual(
            res["Content-Type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("CLIENT BRIEF.docx", res["Content-Disposition"])
        # The bytes must parse as a real Word document and carry the content.
        doc = Document(BytesIO(res.content))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("HFCB CLIENT BRIEF", text)
        self.assertIn("BACKGROUND", text)
        self.assertIn("ii). Staff Loans Scheme", text)
        # Memo table rendered with the client + subject
        self.assertTrue(doc.tables)
        cells = [c.text for row in doc.tables[0].rows for c in row.cells]
        self.assertTrue(any("FROM:" in c for c in cells))
        self.assertTrue(any("SUBJECT:" in c for c in cells))

    def test_requires_authentication(self):
        self.client.force_authenticate(user=None)
        res = self.client.get("/client_briefs/briefs/")
        self.assertIn(res.status_code, (401, 403))

    def test_invalid_opportunities_rejected(self):
        res = self.client.post(
            "/client_briefs/briefs/",
            self._payload(opportunities="not-a-list"),
            format="json",
        )
        self.assertEqual(res.status_code, 400)
